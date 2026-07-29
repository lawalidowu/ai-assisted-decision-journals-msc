#!/usr/bin/env python3
"""Submission build: clean Surrey template + markdown body only (no XML merge).

Usage:
    python scripts/build_submission_docx.py

Output: dissertation/Lawal_MSc_Dissertation.docx

Source of truth for text: dissertation/ABSTRACT.md, CHAPTER_*.md, REFERENCES.md
Template: Official dissertation resources/.../MScDissertationTemplate2026.docx

Does NOT copy formatting from Dissertation_draft.docx — avoids template artefacts,
floating objects, and corrupted styles from prior merges.
"""

from __future__ import annotations

import re
import shutil
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import build_dissertation_docx as md  # noqa: E402

DISS = ROOT / "dissertation"
TEMPLATE = (
    ROOT
    / "Official dissertation resources"
    / "Dissertation templates"
    / "MScDissertationTemplate2026.docx"
)
OUTPUT = DISS / "Lawal_MSc_Dissertation.docx"

TITLE = "AI-assisted decision journaling from public inquiry transcripts: a UK COVID-19 Inquiry case study"
AUTHOR = "Akeeb Idowu Lawal"
SUPERVISOR = "Dr Joaquin Prada"
SUBMISSION_MONTH = "May 2026"
COPYRIGHT_YEAR = "2026"
PATHWAY = "Artificial Intelligence"
DOCUMENT_SUBJECT = "MSc Artificial Intelligence Dissertation"
DOCUMENT_KEYWORDS = (
    "AI-assisted decision journaling; public inquiry transcripts; large language models; UK COVID-19 Inquiry; "
    "decision journal; source traceability; human evaluation"
)

DECLARATION_TEXT = (
    "I confirm that the project dissertation I am submitting is entirely my own work and that "
    "any material used from other sources has been clearly identified and properly acknowledged "
    "and referenced. In submitting this final version of my report to the JISC anti-plagiarism "
    "software resource, I confirm that my work does not contravene the university regulations "
    "on plagiarism as described in the Student Handbook. In so doing I also acknowledge that I "
    "may be held to account for any particular instances of uncited work detected by the JISC "
    "anti-plagiarism software, or as may be found by the project examiner or project organiser. "
    "I also understand that if an allegation of plagiarism is upheld via an Academic "
    "Misconduct Hearing, then I may forfeit any credit for this module or a more severe penalty "
    "may be agreed."
)

INSTRUCTION_LINE = re.compile(r"^<")
SAMPLE_CHAPTER = re.compile(r"^Introduction$", re.I)
TOC_STYLE_PREFIX = "toc "
LEAK_TERMS = (
    "Joint External Evaluation",
    "Jesutomiwa",
    "Kanojia",
    "<Technical CHAPTER>",
    "<Section title>",
    "Example figure and caption",
    "<w:t>Test</w:t>",
    "<w:t>Testing</w:t>",
)


def set_paragraph_text_clean(paragraph, text: str) -> None:
    """Replace entire paragraph content (removes smart tags / extra runs)."""
    element = paragraph._element
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)
    run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    element.append(run)


def set_paragraph_text(paragraph, text: str) -> None:
    set_paragraph_text_clean(paragraph, text)


def set_paragraph_prefix_and_field(paragraph, prefix: str, field_instruction: str) -> None:
    """Paragraph label + Word field (e.g. NUMPAGES)."""
    element = paragraph._element
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)
    prefix_run = OxmlElement("w:r")
    prefix_t = OxmlElement("w:t")
    prefix_t.text = prefix
    prefix_run.append(prefix_t)
    element.append(prefix_run)

    run = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("w:space"), "preserve")
    instr.text = field_instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run.append(fld_begin)
    run.append(instr)
    run.append(fld_sep)
    run.append(fld_end)
    element.append(run)


def count_dissertation_words() -> int:
    paths = [md.ABSTRACT_FILE, *[p for _, p in md.CHAPTER_FILES], md.REFERENCES_FILE]
    total = 0
    for path in paths:
        if not path.exists():
            continue
        in_refs = False
        for line in path.read_text(encoding="utf-8").splitlines():
            stripped = line.strip()
            if path == md.REFERENCES_FILE:
                if stripped.startswith("## References (paste into Word)"):
                    in_refs = True
                    continue
                if in_refs and stripped.startswith("## Verification notes"):
                    break
                if not in_refs:
                    continue
            if stripped.startswith("#") or stripped.startswith("**") or stripped == "---":
                continue
            total += len(re.findall(r"\b[\w'-]+\b", line))
    return total


def patch_word_count_page(doc: Document) -> None:
    words = count_dissertation_words()
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("Number of Pages"):
            set_paragraph_prefix_and_field(p, "Number of Pages:\t", " NUMPAGES ")
        elif text.startswith("Number of Words"):
            set_paragraph_text_clean(p, f"Number of Words:\t{words:,}")


def delete_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def find_heading_index(doc: Document, title: str, style: str = "Unnumbered 1") -> int | None:
    needle = title.strip().lower()
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == style and needle in p.text.strip().lower():
            return i
    return None


def patch_front_matter(doc: Document) -> None:
    paras = doc.paragraphs
    fields = {
        0: TITLE,
        2: AUTHOR,
        5: f"Master of Science in {PATHWAY}",
        17: SUBMISSION_MONTH,
        18: f"Supervised by: {SUPERVISOR}",
        20: f"{AUTHOR}  {COPYRIGHT_YEAR}",
        23: DECLARATION_TEXT,
        27: TITLE,
        31: AUTHOR,
        36: "Author Signature\t_______________________________\tDate: _______________",
        40: f"Supervisor's name: {SUPERVISOR}",
    }
    for idx, text in fields.items():
        if idx < len(paras):
            set_paragraph_text_clean(paras[idx], text)

    # Keep the page-1 Title/Author visual styles, but exclude only these two
    # exact paragraphs from TOCs that collect outline levels 1–3.
    for idx, expected_text in ((0, TITLE), (2, AUTHOR)):
        if idx >= len(paras) or paras[idx].text.strip() != expected_text:
            continue
        p_pr = paras[idx]._element.get_or_add_pPr()
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            p_pr.append(outline)
        outline.set(qn("w:val"), "9")

    # Template splits "from the" / "University of Surrey" across two paragraphs with a smart tag.
    if len(paras) > 6 and "from the" in paras[6].text:
        set_paragraph_text_clean(paras[6], "from the")
    if len(paras) > 7:
        set_paragraph_text_clean(paras[7], "University of Surrey")


def exclude_word_count_from_toc(doc: Document) -> None:
    """Keep the visible template heading but remove heading-linked TOC metadata."""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() != "Word Count":
            continue

        p_pr = paragraph._element.get_or_add_pPr()
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            p_pr.append(outline)
        outline.set(qn("w:val"), "9")

        bookmark_ids = {
            marker.get(qn("w:id"))
            for marker in paragraph._element.findall(qn("w:bookmarkStart"))
            if (marker.get(qn("w:name")) or "").startswith("_Toc")
        }
        for marker in list(paragraph._element):
            if marker.tag in (qn("w:bookmarkStart"), qn("w:bookmarkEnd")):
                if marker.get(qn("w:id")) in bookmark_ids:
                    paragraph._element.remove(marker)

        for run in paragraph.runs:
            r_pr = run._r.get_or_add_rPr()
            r_style = r_pr.find(qn("w:rStyle"))
            if r_style is not None and r_style.get(qn("w:val")) == "Unnumbered1Char":
                r_pr.remove(r_style)
                run.font.name = "Arial"
                run.font.size = Pt(16)
                run.bold = True
                run.font.all_caps = True
        return


def load_abstract_paragraphs() -> list[str]:
    if not md.ABSTRACT_FILE.exists():
        return []
    lines = md.ABSTRACT_FILE.read_text(encoding="utf-8").splitlines()
    paras: list[str] = []
    buf: list[str] = []
    for line in lines:
        if line.strip().startswith("#"):
            continue
        if not line.strip():
            if buf:
                paras.append(" ".join(buf))
                buf = []
            continue
        buf.append(line.strip())
    if buf:
        paras.append(" ".join(buf))
    return paras


def replace_abstract_section(doc: Document) -> None:
    abstract_paras = load_abstract_paragraphs()
    abs_idx = find_heading_index(doc, "Abstract")
    toc_idx = find_heading_index(doc, "Table of Contents")
    if abs_idx is None or toc_idx is None or not abstract_paras:
        return
    for i in range(toc_idx - 1, abs_idx, -1):
        delete_paragraph(doc.paragraphs[i])
    anchor = doc.paragraphs[abs_idx]._element
    prev = anchor
    body_style = md.resolve_style(doc, "Body Text")
    for text in abstract_paras:
        new_p = OxmlElement("w:p")
        prev.addnext(new_p)
        para = Paragraph(new_p, doc.paragraphs[abs_idx]._parent)
        para.style = doc.styles[body_style]
        para.add_run(text)
        prev = new_p


def remove_instruction_paragraphs(doc: Document) -> None:
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        p = doc.paragraphs[i]
        if INSTRUCTION_LINE.match(p.text.strip()):
            delete_paragraph(p)


def delete_toc_snapshot_paragraphs(doc: Document) -> None:
    """Remove static placeholder TOC/LOF lines (not real fields)."""
    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        name = p.style.name
        text = p.text.strip()
        if name.startswith(TOC_STYLE_PREFIX) or name == "table of figures":
            delete_paragraph(p)
            continue
        if "Example figure" in text:
            delete_paragraph(p)
            continue
        if name == "Body First" and text.startswith("<"):
            delete_paragraph(p)
            continue
        i += 1


def delete_sample_body(doc: Document) -> None:
    cut: int | None = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1" and SAMPLE_CHAPTER.match(p.text.strip()):
            cut = i
            break
    if cut is None:
        return
    while cut < len(doc.paragraphs):
        delete_paragraph(doc.paragraphs[cut])


def insert_field_after_heading(doc: Document, heading_title: str, instruction: str) -> None:
    idx = find_heading_index(doc, heading_title)
    if idx is None:
        return
    anchor = doc.paragraphs[idx]._element
    new_p = OxmlElement("w:p")
    anchor.addnext(new_p)
    para = Paragraph(new_p, doc.paragraphs[idx]._parent)
    para.style = doc.styles[md.resolve_style(doc, "Body Text")]
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("w:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def insert_unnumbered_heading_after(doc: Document, after_idx: int, text: str) -> int:
    anchor = doc.paragraphs[after_idx]._element
    new_p = OxmlElement("w:p")
    anchor.addnext(new_p)
    para = Paragraph(new_p, doc.paragraphs[after_idx]._parent)
    para.style = doc.styles["Unnumbered 1"]
    para.add_run(text)
    # Return new index (approximate: after_idx + 1)
    for i, p in enumerate(doc.paragraphs):
        if p._element is new_p:
            return i
    return after_idx + 1


def insert_auto_fields(doc: Document) -> None:
    insert_field_after_heading(doc, "Table of Contents", ' TOC \\o "1-3" \\h \\z \\u ')
    lof_idx = find_heading_index(doc, "List of Figures")
    if lof_idx is None:
        return
    # Native caption lists collect the visible SEQ Figure/Table fields.
    insert_field_after_heading(doc, "List of Figures", ' TOC \\h \\z \\c "Figure" ')
    if find_heading_index(doc, "List of Tables") is None:
        insert_unnumbered_heading_after(doc, lof_idx + 1, "List of Tables")
    insert_field_after_heading(doc, "List of Tables", ' TOC \\h \\z \\c "Table" ')


def patch_headers(doc: Document) -> None:
    for section in doc.sections:
        for p in section.header.paragraphs:
            if "Author Name" in p.text or "AUTHOR" in p.text:
                set_paragraph_text(p, f"{AUTHOR}, MSc dissertation")


def set_section_page_numbering(section, number_format: str, start: int) -> None:
    """Set native Word page-number format and restart value for one section."""
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        insert_before = {
            qn("w:cols"),
            qn("w:formProt"),
            qn("w:vAlign"),
            qn("w:noEndnote"),
            qn("w:titlePg"),
            qn("w:textDirection"),
            qn("w:bidi"),
            qn("w:rtlGutter"),
            qn("w:docGrid"),
            qn("w:printerSettings"),
            qn("w:sectPrChange"),
        }
        index = next(
            (i for i, child in enumerate(sect_pr) if child.tag in insert_before),
            len(sect_pr),
        )
        sect_pr.insert(index, pg_num_type)
    pg_num_type.set(qn("w:fmt"), number_format)
    pg_num_type.set(qn("w:start"), str(start))


def clear_header_footer_content(header_footer) -> Paragraph:
    """Remove inherited/template content and leave one empty paragraph."""
    element = header_footer._element
    for child in list(element):
        element.remove(child)
    return header_footer.add_paragraph()


def append_page_field(paragraph: Paragraph, *, roman: bool = False) -> None:
    """Append an updateable native Word PAGE field.

    When roman=True, use an explicit ``PAGE \\* roman`` switch so lower-case
    Roman numerals render reliably across Word and LibreOffice even when
    section pgNumType alone is not honoured by a viewer.
    """
    begin_run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    begin_run._r.append(fld_begin)

    instruction_run = paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = r" PAGE \* roman " if roman else " PAGE "
    instruction_run._r.append(instruction)

    separator_run = paragraph.add_run()
    fld_separator = OxmlElement("w:fldChar")
    fld_separator.set(qn("w:fldCharType"), "separate")
    separator_run._r.append(fld_separator)

    # Cached display text only; Word/LibreOffice replace on field update.
    paragraph.add_run("ii" if roman else "1")

    end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def configure_page_numbering_sections(doc: Document) -> None:
    """Create Roman front matter and Arabic main-body numbering sections."""
    if len(doc.sections) != 1:
        raise ValueError(
            f"Expected one section after template cleanup, found {len(doc.sections)}"
        )

    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    front_section = doc.sections[0]

    front_section.different_first_page_header_footer = True
    set_section_page_numbering(front_section, "lowerRoman", 1)
    front_section.footer.is_linked_to_previous = False
    front_footer = clear_header_footer_content(front_section.footer)
    front_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    append_page_field(front_footer, roman=True)

    # If odd/even footers are active, keep the same explicit Roman PAGE field.
    if front_section.even_page_footer is not None:
        try:
            front_section.even_page_footer.is_linked_to_previous = False
            even_footer = clear_header_footer_content(front_section.even_page_footer)
            even_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            append_page_field(even_footer, roman=True)
        except Exception:
            pass

    front_section.first_page_footer.is_linked_to_previous = False
    clear_header_footer_content(front_section.first_page_footer)

    main_section.start_type = WD_SECTION.NEW_PAGE
    main_section.different_first_page_header_footer = False
    set_section_page_numbering(main_section, "decimal", 1)
    main_section.footer.is_linked_to_previous = False
    main_footer = clear_header_footer_content(main_section.footer)
    main_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    append_page_field(main_footer, roman=False)


def configure_document_metadata_and_field_updates(doc: Document) -> None:
    """Set submission metadata and ask Word to refresh existing fields on open."""
    properties = doc.core_properties
    properties.title = TITLE
    properties.author = AUTHOR
    properties.subject = DOCUMENT_SUBJECT
    properties.keywords = DOCUMENT_KEYWORDS

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def delete_template_tables(doc: Document) -> None:
    """Remove sample tables left in template body (e.g. Test/More/Testing grid)."""
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:tbl"):
            body.remove(child)


def dedupe_fields_after_headings(doc: Document) -> None:
    for title in ("Table of Contents", "List of Figures", "List of Tables"):
        idx = find_heading_index(doc, title)
        if idx is None:
            continue
        field_indices: list[int] = []
        for j in range(idx + 1, min(idx + 5, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            if p.style.name == "Unnumbered 1":
                break
            if p._element.xpath(".//w:fldChar"):
                field_indices.append(j)
        for j in reversed(field_indices[1:]):
            delete_paragraph(doc.paragraphs[j])


def trim_blank_before_chapter_one(doc: Document) -> None:
    ch1: int | None = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1" and "Chapter 1" in p.text:
            ch1 = i
            break
    if ch1 is None or ch1 == 0:
        return
    for i in range(ch1 - 1, -1, -1):
        p = doc.paragraphs[i]
        if p.style.name == "Unnumbered 1":
            break
        if not p.text.strip() and not p._element.xpath(".//w:drawing"):
            delete_paragraph(p)
        else:
            break


FIGURE_CAPTION_RE = re.compile(r"^Figure\s+[\d.]+(?:a|b)?\.\s*(.+)$", re.I | re.DOTALL)


def finalize_dissertation_in_word(docx_path: Path) -> dict[str, int | bool]:
    """Repaginate, set page count, update TOC/LOF/LOT fields via Word COM."""
    stats: dict[str, int | bool] = {"ok": False, "pages": 0, "figures": 0, "tables": 0}
    try:
        import win32com.client  # type: ignore[import-untyped]
    except ImportError:
        return stats

    word = None
    doc = None
    wdStatisticPages = 2
    wdFieldSequence = 12

    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(docx_path.resolve()), ReadOnly=False, AddToRecentFiles=False)

        doc.Repaginate()
        pages = int(doc.ComputeStatistics(wdStatisticPages))
        stats["pages"] = pages

        for i in range(1, doc.Paragraphs.Count + 1):
            text = doc.Paragraphs(i).Range.Text.strip()
            if text.startswith("Number of Pages"):
                doc.Paragraphs(i).Range.Text = f"Number of Pages:\t{pages}\r"
                break

        # Resolve caption numbers before building their native Figure/Table lists.
        for field in doc.Fields:
            if field.Type == wdFieldSequence:
                field.Update()
        doc.Fields.Update()
        for toc in doc.TablesOfContents:
            toc.Update()
        for tof in doc.TablesOfFigures:
            tof.Update()
        doc.Repaginate()
        doc.Fields.Update()
        for section in doc.Sections:
            for footer_index in (1, 2, 3):
                footer = section.Footers(footer_index)
                if footer.Exists:
                    footer.Range.Fields.Update()

        stats["figures"] = sum(
            1
            for field in doc.Fields
            if field.Type == wdFieldSequence and field.Code.Text.strip().startswith("SEQ Figure")
        )
        stats["tables"] = sum(
            1
            for field in doc.Fields
            if field.Type == wdFieldSequence and field.Code.Text.strip().startswith("SEQ Table")
        )

        doc.Save()
        stats["ok"] = True
        return stats
    except Exception as exc:
        stats["error"] = str(exc)  # type: ignore[assignment]
        return stats
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def patch_header_footer_xml(docx_path: Path) -> None:
    tmp = docx_path.with_suffix(".patch.tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/header") and item.filename.endswith(".xml"):
                text = data.decode("utf-8")
                text = text.replace("Author Name", AUTHOR)
                text = text.replace("Jesutomiwa Salam", AUTHOR)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(docx_path)


def disable_reference_style_numbering(doc: Document) -> None:
    """Remove automatic list numbering from the Reference paragraph style."""
    try:
        style = doc.styles["Reference"]
    except KeyError:
        return
    p_pr = style._element.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn("w:numPr"):
            p_pr.remove(child)


def add_reference_paragraph(doc: Document, ref: str, style: str, number: int) -> None:
    """IEEE reference with explicit '[n] ' spacing and hanging indent."""
    style_name = md.resolve_style(doc, style)
    p = doc.add_paragraph(style=style_name)
    md.clear_paragraph_numbering(p)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    pf.space_after = Pt(6)
    marker = p.add_run(f"[{number}] ")
    marker.font.name = "Times New Roman"
    marker.font.size = Pt(12)
    for part in re.split(r"(\*[^*]+\*)", ref):
        if not part:
            continue
        if part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.italic = True
        else:
            run = p.add_run(part)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def append_body_from_markdown(doc: Document) -> None:
    body_style = md.resolve_style(doc, "Body Text")
    caption_style = md.resolve_style(doc, "Figure")
    ref_style = md.resolve_style(doc, "Reference")

    for chapter_index, (chapter_title, md_path) in enumerate(md.CHAPTER_FILES):
        md.add_heading(doc, chapter_title, level=1)
        if chapter_index > 0:
            doc.paragraphs[-1].paragraph_format.page_break_before = True
        md.render_markdown(
            doc,
            md_path,
            body_style=body_style,
            caption_style=caption_style,
        )

    refs = md.load_references()
    if refs:
        # Heading 1 already starts a new page in the Surrey template. Do not also
        # insert an explicit PAGE break, or a blank page appears before References.
        disable_reference_style_numbering(doc)
        md.add_heading(doc, "References", level=1)
        doc.paragraphs[-1].paragraph_format.page_break_before = True
        for i, ref in enumerate(refs, start=1):
            add_reference_paragraph(doc, ref, ref_style, i)


def verify_output(docx_path: Path) -> list[str]:
    issues: list[str] = []
    with zipfile.ZipFile(docx_path) as z:
        parts = [z.read(n).decode("utf-8", errors="ignore") for n in z.namelist() if n.endswith(".xml")]
        blob = "\n".join(parts)
        for term in LEAK_TERMS:
            if term in blob:
                issues.append(term)
        if "Opening" in blob and ">Opening<" in blob.replace(" ", ""):
            issues.append("Opening heading leaked")
    doc = Document(docx_path)
    for p in doc.paragraphs:
        if p.text.strip() == "Opening":
            issues.append("Visible 'Opening' paragraph")
    return issues


def prepare_template_copy() -> Path:
    candidates = [
        OUTPUT,
        DISS / "Lawal_MSc_Dissertation_submission.docx",
    ]
    for candidate in candidates:
        try:
            shutil.copy2(TEMPLATE, candidate)
            if candidate != OUTPUT:
                print(f"Note: {OUTPUT.name} is open - wrote {candidate.name}. Close Word and re-run.")
            return candidate
        except PermissionError:
            continue
    raise SystemExit("Close all Lawal_MSc_Dissertation*.docx files in Word and re-run.")


def try_update_fields_in_word(docx_path: Path) -> bool:
    """Legacy wrapper — use finalize_dissertation_in_word."""
    return bool(finalize_dissertation_in_word(docx_path).get("ok"))


def build(*, word_finalize: bool = False) -> Path:
    if not TEMPLATE.is_file():
        raise SystemExit(f"Template not found: {TEMPLATE}")

    out = prepare_template_copy()
    doc = Document(out)

    md.clear_style_numbering(doc)
    md.disable_document_hyphenation(doc)
    patch_front_matter(doc)
    exclude_word_count_from_toc(doc)
    replace_abstract_section(doc)
    remove_instruction_paragraphs(doc)
    delete_toc_snapshot_paragraphs(doc)
    delete_sample_body(doc)
    delete_template_tables(doc)

    insert_auto_fields(doc)
    dedupe_fields_after_headings(doc)
    trim_blank_before_chapter_one(doc)
    configure_page_numbering_sections(doc)
    append_body_from_markdown(doc)
    patch_word_count_page(doc)
    patch_headers(doc)
    configure_document_metadata_and_field_updates(doc)

    doc.save(out)
    patch_header_footer_xml(out)

    finalize = finalize_dissertation_in_word(out) if word_finalize else {"ok": False}
    if finalize.get("ok"):
        print(
            f"Word finalize: {finalize.get('pages', '?')} pages, "
            f"{finalize.get('figures', 0)} figure captions, "
            f"{finalize.get('tables', 0)} table captions."
        )
    elif word_finalize:
        raise SystemExit(
            f"Native Word field finalization failed: {finalize.get('error', 'unknown error')}. "
            "Submission DOCX is not ready."
        )
    else:
        print("Native Word field finalization was explicitly skipped.")

    issues = verify_output(out)
    if issues:
        print("WARNING - verify in Word:", ", ".join(issues))

    size_kb = out.stat().st_size // 1024
    print(f"Wrote {out} ({size_kb} KB)")
    print("Submission build: markdown -> Surrey template (no draft XML merge)")
    print(f"Word count (body, from markdown): {count_dissertation_words():,}")
    print("In Word:")
    print("  1. Verify TOC, List of Figures, and List of Tables")
    print("  2. Sign declaration; confirm page/word counts")
    print("  3. Section breaks + roman/Arabic page numbers")
    print("  4. Read PDF once (SUBMISSION_CHECKLIST.md)")
    return out


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Build Surrey submission dissertation.")
    parser.add_argument(
        "--no-word-finalize",
        action="store_true",
        help="Skip native Word field updates after building (not recommended for submission output).",
    )
    args = parser.parse_args()
    build(word_finalize=not args.no_word_finalize)
