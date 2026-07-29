#!/usr/bin/env python3
"""Build dissertation Word draft from markdown chapter files.

Usage:
    python scripts/build_dissertation_docx.py
    python scripts/build_dissertation_docx.py --output dissertation/Dissertation_draft.docx

Reads markdown from dissertation/CHAPTER_*.md and writes a .docx with Heading 1/2 styles.
Strips markdown metadata blocks (status lines, horizontal rules used as separators).
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING

HEADING_STYLES = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
}
SKIP_SECTION_HEADINGS = {"Opening"}  # editorial section labels — not printed in Word
NON_TOC_CONTENT_LABELS = frozenset(
    {
        "Aim",
        "Objectives",
        "Project evolution",
        "Phase 1 (complete)",
        "Phase 2 (2a–2c complete)",
        "Principal artefact",
        "Explicitly out of scope",
        "Rubric A: Valid decision journal entry?",
        "Rubric B: Strength of support",
    }
)

ROOT = Path(__file__).resolve().parents[1]
DISS = ROOT / "dissertation"

CHAPTER_FILES = [
    ("Chapter 1 — Introduction", DISS / "CHAPTER_1_INTRODUCTION.md"),
    ("Chapter 2 — Literature review", DISS / "CHAPTER_2_LITERATURE.md"),
    ("Chapter 3 — Methodology", DISS / "CHAPTER_3_METHODS.md"),
    ("Chapter 4 — Results and evaluation", DISS / "CHAPTER_4_RESULTS.md"),
    ("Chapter 5 — Discussion and conclusion", DISS / "CHAPTER_5_DISCUSSION.md"),
    ("Appendix A — Manual excerpt evidence", DISS / "APPENDIX_A_MANUAL_EXCERPTS.md"),
    ("Appendix B — Supporting extraction and robustness results", DISS / "APPENDIX_B_SUPPORTING_RESULTS.md"),
    ("Appendix C — Clustering supplement", DISS / "APPENDIX_C_CLUSTERING_SUPPLEMENT.md"),
]

ABSTRACT_FILE = DISS / "ABSTRACT.md"
REFERENCES_FILE = DISS / "REFERENCES.md"

# Explicit figure markers in markdown — insert image + one caption only (no duplicate body text)
FIGURE_MARKERS: dict[str, tuple[str, str]] = {
    "[[FIGURE:3.1]]": (
        "outputs/figures/conceptual_framework.png",
        "Figure 3.1. Two-phase workflow from public inquiry transcripts to a fixed reference dataset of candidate decision-journal entries. Phase 1 covers extraction and source traceability. Phase 2 applies review flags, n = 50 human validation, exploratory clustering for navigation, and a supplementary n = 60 Joint External Evaluation and Decision Quality mapping pilot. Human judgement remains the final authority for interpretation.",
    ),
    "[[FIGURE:3.2]]": (
        "outputs/figures/implemented_pipeline.png",
        "Figure 3.2. Technical implementation pipeline from eight inquiry transcripts to the fixed reference dataset and the separate evaluation, review, clustering and framework-mapping activities.",
    ),
    "[[FIGURE:4.3]]": (
        "outputs/figures/error_taxonomy_distribution.png",
        "Figure 4.3. Assigned error categories in the stratified sample (n = 42).",
    ),
    "[[FIGURE:4.9]]": (
        "outputs/figures/figure4_9_rubric_crosstab.png",
        "Figure 4.9. Rubric A × Rubric B cross-tabulation (n = 50).",
    ),
    "[[FIGURE:4.10]]": (
        "outputs/figures/phase1_cluster_sizes.png",
        "Figure 4.10. Theme distribution across 20 heuristic groups.",
    ),
}

TABLE_CAPTIONS: dict[str, tuple[str, ...]] = {
    # Wave 5B: Table 1.1 removed with Section 1.4 scope rewrite.
    "CHAPTER_1_INTRODUCTION.md": (),
    "CHAPTER_2_LITERATURE.md": (),
    "CHAPTER_3_METHODS.md": (
        "Table 3.1. Evaluation methods and the populations or samples they address.",
        "Table 3.2. Phase 1 hearing transcript sample.",
        "Table 3.3. Decision object schema.",
        "Table 3.4. Non-destructive review-flag rules.",
        "Table 3.5. Stratified confidence-validation sample (n = 50).",
        "Table 3.6. Rubric A journal-validity criteria.",
        "Table 3.7. Rubric B evidence-strength criteria.",
        "Table 3.8. Rubric A × Rubric B interpretation framework.",
        "Table 3.9. Decision Quality elements used in the n = 60 mapping pilot.",
        "Table 3.10. Items deferred beyond MSc scope.",
    ),
    "CHAPTER_4_RESULTS.md": (
        "Table 4.1. Phase 1 corpus totals.",
        "Table 4.2. Manual triangulation results.",
        "Table 4.3. Automated confidence agreement with human Rubric B.",
        "Table 4.4. Framework-mapping outcomes in the n = 60 pilot.",
        "Table 4.5. Candidate-statement faithfulness classifications (n = 60).",
        "Table 4.6. Supplementary robustness and transfer checks.",
    ),
    "APPENDIX_A_MANUAL_EXCERPTS.md": (
        "Table A.1. Manual excerpt index for triangulation.",
    ),
    "APPENDIX_B_SUPPORTING_RESULTS.md": (
        "Table B.1. Per-hearing extraction and traceability results.",
        "Table B.2. Default and inquiry prompt extraction counts.",
        "Table B.3. Error taxonomy definitions and counts (n = 42).",
        "Table B.4. Keyword-baseline results by excerpt.",
        "Table B.5. GRACE-adapted quality scores.",
        "Table B.6. Report-genre pilot summary.",
        "Table B.7. Structural reliability stress-test results.",
        "Table B.8. n = 60 review provenance routes.",
    ),
    "APPENDIX_C_CLUSTERING_SUPPLEMENT.md": (
        "Table C.1. Largest heuristic themes in the candidate corpus.",
        "Table C.2. Additional mid-sized heuristic themes.",
    ),
    # Wave 4B: no Chapter 5 tables (argumentative summary tables removed).
    "CHAPTER_5_DISCUSSION.md": (),
}

SKIP_LINE_PATTERNS = [
    re.compile(r"^\*\*Status:\*\*"),
    re.compile(r"^\*\*Title \(suggested\):\*\*"),
    re.compile(r"^---$"),
    re.compile(r"^\*\*Figure references:\*\*"),
    re.compile(r"^# Chapter \d+"),  # duplicate top heading — we add chapter title separately
    re.compile(r"^# Appendix [A-Z]"),  # duplicate top heading — we add appendix title separately
    re.compile(r"^\*\*Note:\*\* Do not"),
    re.compile(r"^\*\*Title:\*\*"),
    re.compile(r"^# Abstract$"),
    re.compile(r"^\*.*(?:Draft|Sources:|Paste into).*\*$"),
]


def should_skip(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.startswith("## Draft checklist"):
        return True
    return any(p.match(stripped) for p in SKIP_LINE_PATTERNS)


def clear_paragraph_numbering(paragraph) -> None:
    """Remove Word multilevel list numbering from a paragraph (Surrey template conflict)."""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


def format_paragraph(paragraph, *, heading: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.widow_control = True
    if heading:
        clear_paragraph_numbering(paragraph)


def format_caption(paragraph, *, figure: bool = False) -> None:
    """Apply consistent black Times New Roman caption formatting (not Heading blue)."""
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.bold = False
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.widow_control = True
    if figure:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def disable_document_hyphenation(doc: Document) -> None:
    """Prevent mid-word hyphenation that breaks table labels and headings."""
    settings = doc.settings.element
    for tag in ("w:autoHyphenation", "w:doNotHyphenateCaps"):
        # Remove existing auto-hyphenation enablement if present.
        node = settings.find(qn(tag))
        if node is not None and tag == "w:autoHyphenation":
            settings.remove(node)
    auto = OxmlElement("w:autoHyphenation")
    auto.set(qn("w:val"), "0")
    settings.append(auto)


def strip_markdown_emphasis(text: str) -> str:
    """Remove markdown emphasis markers for plain dissertation prose."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


# Exact reader-facing text for Section 3.2.3 operational-definition callout.
SECTION_323_OPERATIONAL_DEFINITION_PLAIN = (
    "Decision journal entries record pandemic-response agreements, adopted measures, "
    "and authoritative directions (who decided or agreed what), as stated in the source "
    "text, including when recalled in inquiry evidence, each backed by a verbatim source "
    "quote. Witness opinion, speculation, and options that were not agreed are excluded."
)

# Authorised bold-italic phrases only (Section 3.2.3).
SECTION_323_BOLD_ITALIC_PHRASES = (
    "Decision journal entries",
    "agreements, adopted measures, and authoritative directions",
    "verbatim source quote",
)


def is_section_323_operational_definition_blockquote(quote: str) -> bool:
    """True only for the Section 3.2.3 operational-definition callout."""
    plain = strip_markdown_emphasis(quote).strip()
    if plain == SECTION_323_OPERATIONAL_DEFINITION_PLAIN:
        return True
    # Anchor on unique Markdown prefix as a secondary guard.
    return quote.strip().startswith("**Decision journal entries**")


def add_section_323_operational_definition_runs(paragraph, quote: str) -> None:
    """Targeted rich-run construction for Section 3.2.3 only (not global Markdown)."""
    plain = strip_markdown_emphasis(quote).strip()
    if plain != SECTION_323_OPERATIONAL_DEFINITION_PLAIN:
        # Safety: never invent wording; fall back to stripped italic plain text.
        run = paragraph.add_run(plain)
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        return

    # Build runs from the authorised plain text, bold-italic only on the three phrases.
    remaining = plain
    for phrase in SECTION_323_BOLD_ITALIC_PHRASES:
        idx = remaining.find(phrase)
        if idx < 0:
            continue
        if idx > 0:
            run = paragraph.add_run(remaining[:idx])
            run.bold = False
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        run = paragraph.add_run(phrase)
        run.bold = True
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        remaining = remaining[idx + len(phrase) :]
    if remaining:
        run = paragraph.add_run(remaining)
        run.bold = False
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_runs_plain(paragraph, text: str) -> None:
    """Body text: strip Markdown emphasis markers; do not apply global bold."""
    run = paragraph.add_run(strip_markdown_emphasis(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def resolve_style(doc: Document, name: str, fallback: str = "Normal") -> str:
    try:
        _ = doc.styles[name]
        return name
    except KeyError:
        return fallback


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    style_name = resolve_style(doc, style or "Normal")
    p = doc.add_paragraph(style=style_name)
    add_runs_plain(p, text)
    is_heading = bool(style_name.startswith("Heading"))
    format_paragraph(p, heading=is_heading)


def add_content_label(doc: Document, text: str) -> None:
    """Render an exact non-TOC content label with Heading 3-like emphasis."""
    style = resolve_style(doc, "Body Text", "Normal")
    p = doc.add_paragraph(style=style)
    add_runs_plain(p, text)
    for run in p.runs:
        run.bold = True
    format_paragraph(p)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12)
    p_pr = p._element.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), "9")


def add_heading(doc: Document, text: str, level: int) -> None:
    """Add a heading with manual section numbers only (no Word auto-numbering)."""
    if text in NON_TOC_CONTENT_LABELS:
        add_content_label(doc, text)
        return
    style = resolve_style(doc, HEADING_STYLES.get(level, "Heading 2"))
    p = doc.add_paragraph(style=style)
    add_runs_plain(p, text)
    format_paragraph(p, heading=True)
    clear_paragraph_numbering(p)


def parse_table_lines(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if not line.strip().startswith("|"):
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        rows.append(cells)
    return rows


GENERIC_TABLE_HEADERS = frozenset(
    {
        "metric", "value", "step", "status", "outcome", "evidence", "direction",
        "finding", "layer", "question", "result", "priority", "rationale",
        "hearing date", "chunks", "decisions", "error type", "pass rate",
        "why beyond msc scope", "phase", "alignment", "signal",
    }
)


def strip_section_number(title: str) -> str:
    return re.sub(r"^[\d.]+\s*", "", title).strip()


def infer_table_caption(rows: list[list[str]], section_title: str | None) -> str:
    section = strip_section_number(section_title) if section_title else ""
    if section:
        return section[:90]
    if not rows or not rows[0]:
        return "Summary"
    header = strip_markdown_emphasis(rows[0][0]).strip()
    if header.lower() not in GENERIC_TABLE_HEADERS:
        return header[:90]
    if len(rows) > 1 and rows[1][0]:
        return strip_markdown_emphasis(rows[1][0]).strip()[:90]
    return header[:90]


def clear_paragraph_numbering(paragraph) -> None:
    """Remove style-based auto-numbering so explicit captions are not prefixed."""
    p_pr = paragraph._element.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn("w:numPr"):
            p_pr.remove(child)


def append_seq_field(paragraph, identifier: str, value: int) -> None:
    """Append a visible native Word SEQ field reset to the supplied value."""
    begin_run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    begin_run._r.append(fld_begin)

    instruction_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("w:space"), "preserve")
    instr.text = f" SEQ {identifier} \\* ARABIC \\r {value} "
    instruction_run._r.append(instr)

    separator_run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    separator_run._r.append(fld_sep)

    paragraph.add_run(str(value))

    end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def append_native_caption(paragraph, caption_text: str, identifier: str) -> None:
    """Write the approved caption text using a native SEQ Figure/Table field."""
    pattern = rf"^{identifier} (\d+)\.(\d+)([a-z]?)\. (.+)$"
    match = re.fullmatch(pattern, caption_text)
    if match is not None:
        chapter, number, suffix, title = match.groups()
        paragraph.add_run(f"{identifier} {chapter}.")
        append_seq_field(paragraph, identifier, int(number))
        paragraph.add_run(f"{suffix}. {title}")
        return
    # Appendix captions use letter prefixes (Table A.1 / Figure A.1) and are written
    # as fixed text so they do not enter the chapter SEQ stream.
    appendix_pattern = rf"^{identifier} ([A-Z])\.(\d+)([a-z]?)\. (.+)$"
    appendix_match = re.fullmatch(appendix_pattern, caption_text)
    if appendix_match is not None:
        chapter, number, suffix, title = appendix_match.groups()
        paragraph.add_run(f"{identifier} {chapter}.{number}{suffix}. {title}")
        return
    raise ValueError(f"Invalid {identifier} caption format: {caption_text}")


def add_markdown_table(
    doc: Document,
    rows: list[list[str]],
    *,
    section_title: str | None = None,
    caption_text: str | None = None,
) -> None:
    if not rows:
        return
    cap_style = resolve_style(doc, "Caption", "Normal")
    header_label = caption_text or infer_table_caption(rows, section_title)
    cap = doc.add_paragraph(style=cap_style)
    cap.paragraph_format.keep_with_next = True
    keep_whole_tables = {
        "Table 3.1. Evaluation methods and the populations or samples they address.",
        "Table 3.9. Decision Quality elements used in the n = 60 mapping pilot.",
        "Table 3.10. Items deferred beyond MSc scope.",
    }
    keep_whole_table = header_label in keep_whole_tables
    left_align_table = bool(header_label) and (
        header_label.startswith("Table 3.")
        or header_label.startswith("Table 4.")
        or header_label.startswith("Table A.")
        or header_label.startswith("Table B.")
        or header_label.startswith("Table C.")
    )
    # Table 3.1 follows a large figure: force a clean page start and keep the
    # compact methods table together. Chapter 4 tables must NOT take an
    # unconditional page break; cantSplit alone prevents mid-row splits while
    # still allowing a small table to begin under its introductory prose.
    if keep_whole_table:
        cap.paragraph_format.page_break_before = True
    clear_paragraph_numbering(cap)
    append_native_caption(cap, header_label, "Table")
    format_caption(cap, figure=False)

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    last_row_idx = len(table.rows) - 1
    compact = bool(header_label) and (
        header_label.startswith("Table A.")
        or header_label.startswith("Table B.")
        or header_label.startswith("Table C.")
        or header_label.startswith("Table 4.")
    )
    for i, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        # Repeat header row on continuation pages.
        if i == 0 and tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))

    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < ncols:
                cell_obj = table.rows[i].cells[j]
                cell_obj.text = strip_markdown_emphasis(cell)
                for paragraph in cell_obj.paragraphs:
                    if left_align_table or keep_whole_table:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.widow_control = True
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        if compact:
                            run.font.size = Pt(10 if header_label.startswith("Table 4.") else 9)
                        else:
                            run.font.size = Pt(11)
                    # Keep compact methods tables together where practicable.
                    if keep_whole_table and i < last_row_idx:
                        paragraph.paragraph_format.keep_with_next = True
    doc.add_paragraph()


def add_figure(
    doc: Document,
    image_path: Path,
    caption: str,
    width_inches: float = 5.5,
    *,
    caption_style: str = "Normal",
) -> bool:
    """Insert centred figure + explicit caption text (no SEQ auto-number). Returns False if image missing."""
    body_style = resolve_style(doc, "Body Text", "Normal")
    # The Surrey template's Figure style inherits automatic list numbering
    # ("Figure %1."). Use its non-numbered Caption style because the visible
    # caption already carries the intended chapter-based number.
    cap_style = resolve_style(doc, "Caption", "Normal")
    if not image_path.is_file():
        add_paragraph(
            doc,
            f"[Figure file missing: {image_path.name}. Run figure build scripts.]",
            style=body_style,
        )
        return False
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.keep_with_next = True
    pic_para.add_run().add_picture(str(image_path), width=Inches(width_inches))
    visible = strip_markdown_emphasis(caption)
    cap = doc.add_paragraph(style=cap_style)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_paragraph_numbering(cap)
    append_native_caption(cap, visible, "Figure")
    format_caption(cap, figure=True)
    doc.add_paragraph()
    return True


def insert_figure_marker(doc: Document, marker: str, *, caption_style: str = "Normal") -> None:
    if marker not in FIGURE_MARKERS:
        add_paragraph(doc, f"[Unknown figure marker: {marker}]", style="Normal")
        return
    rel_path, caption = FIGURE_MARKERS[marker]
    add_figure(doc, ROOT / rel_path, caption, caption_style=caption_style)


def render_markdown(
    doc: Document,
    md_path: Path,
    *,
    body_style: str = "Normal",
    caption_style: str = "Normal",
) -> None:
    if not md_path.exists():
        add_paragraph(doc, f"[Missing: {md_path.name}]", style="Normal")
        return

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    i = 0
    in_checklist = False
    current_h2: str | None = None
    current_h3: str | None = None
    table_ordinal = 0
    expected_table_captions = TABLE_CAPTIONS.get(md_path.name)

    while i < len(lines):
        line = lines[i]

        if should_skip(line):
            i += 1
            continue

        if line.strip().startswith("## Draft checklist"):
            in_checklist = True
            i += 1
            continue

        if in_checklist:
            i += 1
            continue

        stripped = line.strip()
        if stripped in FIGURE_MARKERS:
            insert_figure_marker(doc, stripped, caption_style=caption_style)
            i += 1
            continue

        # Fenced code blocks (ASCII diagrams) — skip; figures replace these in Word
        if line.strip().startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1  # closing fence
            continue

        # Markdown table block
        if line.strip().startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            caption_text = None
            if expected_table_captions is not None:
                if table_ordinal >= len(expected_table_captions):
                    raise ValueError(
                        f"Unexpected table {table_ordinal + 1} in {md_path.name}; "
                        f"caption map defines {len(expected_table_captions)} tables"
                    )
                caption_text = expected_table_captions[table_ordinal]
            table_ordinal += 1
            add_markdown_table(
                doc,
                parse_table_lines(table_lines),
                section_title=current_h3 or current_h2,
                caption_text=caption_text,
            )
            continue

        # Headings — check #### before ### (four-hash headings must not fall through as body text)
        if line.startswith("#### "):
            title = line[5:].strip()
            current_h3 = title
            add_heading(doc, title, level=3)
            i += 1
            continue
        if line.startswith("### "):
            title = line[4:].strip()
            current_h3 = title
            add_heading(doc, title, level=3)
            i += 1
            continue
        if line.startswith("## "):
            title = line[3:].strip()
            if title in SKIP_SECTION_HEADINGS:
                i += 1
                continue
            current_h2 = title
            current_h3 = None
            add_heading(doc, title, level=2)
            i += 1
            continue

        # Blockquote — italic callouts / Appendix A transcript passages
        if line.startswith("> "):
            quote = line[2:].strip()
            # Collect continued blockquote lines
            while i + 1 < len(lines) and lines[i + 1].startswith("> "):
                i += 1
                quote += " " + lines[i][2:].strip()
            p = doc.add_paragraph(style=body_style)
            if is_section_323_operational_definition_blockquote(quote):
                # Targeted Section 3.2.3 emphasis only (no global Markdown rewrite).
                add_section_323_operational_definition_runs(p, quote)
            else:
                # Pre-micro-closeout behaviour: one italic run; markers stripped.
                # format_paragraph below restores Times New Roman 12 pt (Appendix A).
                run = p.add_run(strip_markdown_emphasis(quote))
                run.italic = True
            format_paragraph(p)
            # Re-apply italic (and authorised bold) after format_paragraph size reset.
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # List items (plain Normal; strip markdown; do not use lstrip on '*' — it breaks **bold**)
        if re.match(r"^[-*]\s+", line.strip()):
            content = re.sub(r"^[-*]\s+", "", line.strip())
            add_paragraph(doc, content, style=body_style)
            i += 1
            continue
        # Numbered lists: keep manual numbers in plain text (avoid Word continuing 1,2,3 → 6,7,8)
        if re.match(r"^\d+\.\s+", line.strip()):
            add_paragraph(doc, line.strip(), style=body_style)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        add_paragraph(doc, stripped, style=body_style)
        i += 1

    if expected_table_captions is not None and table_ordinal != len(expected_table_captions):
        raise ValueError(
            f"Table count mismatch in {md_path.name}: rendered {table_ordinal}, "
            f"caption map defines {len(expected_table_captions)}"
        )


def load_references() -> list[str]:
    """Return IEEE reference lines from REFERENCES.md (between heading and verification notes)."""
    if not REFERENCES_FILE.exists():
        return []
    text = REFERENCES_FILE.read_text(encoding="utf-8")
    lines = text.splitlines()
    refs: list[str] = []
    in_refs = False
    for line in lines:
        if line.strip().startswith("## References (paste into Word)"):
            in_refs = True
            continue
        if in_refs and line.strip().startswith("## Verification notes"):
            break
        if in_refs and line.strip() and not line.strip().startswith("#"):
            if line.strip() in ("---", "—"):
                continue
            refs.append(line.strip())
    return refs


def clear_style_numbering(doc: Document) -> None:
    """Remove multilevel list numbering from Heading 1–3 style definitions (Surrey template)."""
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        pPr = style._element.find(qn("w:pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)


def build_docx(output: Path) -> None:
    doc = Document()
    clear_style_numbering(doc)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    add_paragraph(doc, "Dissertation draft — AI-assisted decision journaling from public inquiry transcripts: a UK COVID-19 Inquiry case study", style="Title")
    add_paragraph(doc, "[Paste Surrey title page and declaration from sample PDF]", style="Normal")
    doc.add_page_break()

    add_heading(doc, "Abstract", level=1)
    render_markdown(doc, ABSTRACT_FILE)
    doc.add_page_break()

    add_paragraph(doc, "[Insert table of contents, list of figures, list of tables — Word auto-generate]", style="Normal")
    doc.add_page_break()

    for chapter_title, md_path in CHAPTER_FILES:
        add_heading(doc, chapter_title, level=1)
        render_markdown(doc, md_path)
        doc.add_page_break()

    refs = load_references()
    if refs:
        add_heading(doc, "References", level=1)
        for ref in refs:
            clean = re.sub(r"\*(.+?)\*", r"\1", ref)
            add_paragraph(doc, clean, style="Normal")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(f"Wrote {output} ({output.stat().st_size // 1024} KB)")


def main() -> None:
    parser = argparse.ArgumentParser(description="Build dissertation Word draft from markdown chapters.")
    parser.add_argument(
        "--output",
        type=Path,
        default=DISS / "Dissertation_draft.docx",
        help="Output .docx path",
    )
    args = parser.parse_args()
    build_docx(args.output.resolve())


if __name__ == "__main__":
    main()
