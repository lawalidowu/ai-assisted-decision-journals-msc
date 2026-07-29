#!/usr/bin/env python3
import zipfile
from pathlib import Path
from docx import Document

path = Path(r"c:\SURREY\MODULES\SEMESTER 2\MSC PROJECT\code\dissertation\Lawal_MSc_Dissertation_build.docx")
doc = Document(path)
print(f"Size: {path.stat().st_size // 1024} KB")
print(f"Paragraphs: {len(doc.paragraphs)}")

with zipfile.ZipFile(path) as z:
    media = [n for n in z.namelist() if n.startswith("word/media/")]
    print(f"Media files: {len(media)}")
    xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
    for term in ("Joint External Evaluation", "Jesutomiwa", "Kanojia", "Example figure", "414"):
        print(f"  leak {term!r}: {term in xml}")
    for name in z.namelist():
        if name.startswith("word/header") and name.endswith(".xml"):
            h = z.read(name).decode("utf-8", errors="ignore")
            print(f"  header {name}: Akeeb={('Akeeb' in h)}, placeholder={('Author Name' in h)}")

abs_i = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Abstract"), None)
if abs_i is not None:
    print("Abstract section:")
    for j in range(abs_i, min(abs_i + 6, len(doc.paragraphs))):
        t = doc.paragraphs[j].text.strip()[:90]
        print(f"  {j} [{doc.paragraphs[j].style.name}]: {t}")

for i, p in enumerate(doc.paragraphs):
    if "Chapter 1" in p.text:
        print(f"Chapter 1 at para {i}")
        break
