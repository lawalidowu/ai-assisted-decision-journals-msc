#!/usr/bin/env python3
"""Wave 7A: rebuild dissertation after September 2026 title-month gate and package FINAL binaries."""
from __future__ import annotations

import hashlib
import json
import re
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))
from build_submission_docx import SUBMISSION_MONTH, count_dissertation_words  # noqa: E402

DOCX = ROOT / "dissertation" / "Lawal_MSc_Dissertation.docx"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


def main() -> None:
    if SUBMISSION_MONTH != "September 2026":
        raise SystemExit(f"SUBMISSION_MONTH is {SUBMISSION_MONTH!r}; expected 'September 2026'")

    print("1) build_submission_docx.py", flush=True)
    subprocess.check_call([sys.executable, str(ROOT / "scripts" / "build_submission_docx.py")], cwd=ROOT)

    print("2) finalize_word_fields.py", flush=True)
    subprocess.check_call([sys.executable, str(ROOT / "scripts" / "finalize_word_fields.py")], cwd=ROOT)

    md_words = count_dissertation_words()
    pages_line = words_line = None
    title_month_ok = False
    sig_blank = False
    for p in Document(str(DOCX)).paragraphs:
        t = p.text.strip()
        if t.startswith("Number of Pages"):
            pages_line = t
        if t.startswith("Number of Words"):
            words_line = t
        if t == "September 2026":
            title_month_ok = True
        if t == "May 2026":
            raise SystemExit("Active DOCX still contains standalone title-month paragraph 'May 2026'")
        if "Author Signature" in t or t.startswith("Author Signature"):
            # signature line should not contain a fabricated name after the label alone is OK
            pass

    blob = "\n".join(p.text for p in Document(str(DOCX)).paragraphs)
    if "September 2026" not in blob:
        raise SystemExit("September 2026 not found in rebuilt DOCX")
    # Declaration signature/date should remain instructional blanks
    if re.search(r"Author Signature\s*[A-Za-z]{2,}", blob) and "Author Signature" in blob:
        # template uses 'Author Signature' label; reject if a signed name is glued
        after = blob.split("Author Signature", 1)[1][:80]
        if re.search(r"^\s*[A-Za-z]{3,}", after) and "Date" not in after[:20]:
            raise SystemExit(f"Possible fabricated signature: {after!r}")
    sig_blank = "Author Signature" in blob and "Date" in blob

    if not words_line or not pages_line:
        raise SystemExit("Missing Number of Words/Pages fields in DOCX")

    disp_words = int(re.sub(r"[^0-9]", "", words_line))
    disp_pages = int(re.sub(r"[^0-9]", "", pages_line))
    print("MD_WORDS", md_words)
    print("DISPLAYED_WORDS", disp_words)
    print("DISPLAYED_PAGES", disp_pages)
    print("TITLE_MONTH_OK", title_month_ok or ("September 2026" in blob))
    print("SIG_FIELDS_PRESENT", sig_blank)
    if disp_words != md_words:
        raise SystemExit(f"DISCREPANCY words md={md_words} displayed={disp_words}")

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    pkg = ROOT / "outputs" / "dissertation_integration" / f"run_{stamp}_wave7a_title_page_september"
    pkg.mkdir(parents=True, exist_ok=False)
    print("PACKAGE", pkg, flush=True)

    import win32com.client  # type: ignore

    wd_format_pdf = 17
    wd_statistic_pages = 2
    pdf_path = pkg / "Lawal_Akeeb_Idowu_MSc_Dissertation_FINAL.pdf"
    docx_out = pkg / "Lawal_Akeeb_Idowu_MSc_Dissertation_FINAL.docx"
    shutil.copy2(DOCX, docx_out)

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    doc = None
    try:
        doc = word.Documents.Open(str(docx_out.resolve()), ReadOnly=True, AddToRecentFiles=False)
        doc.Repaginate()
        physical_pages = int(doc.ComputeStatistics(wd_statistic_pages))
        doc.ExportAsFixedFormat(str(pdf_path.resolve()), wd_format_pdf)
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()

    docx_hash = sha256(docx_out)
    pdf_hash = sha256(pdf_path)

    # PDF text check
    import fitz

    pdf_text = "\n".join(page.get_text("text") or "" for page in fitz.open(pdf_path))
    if "September 2026" not in pdf_text:
        raise SystemExit("PDF missing September 2026")
    if re.search(r"(?m)^May 2026\s*$", pdf_text):
        raise SystemExit("PDF still has title-month line May 2026")

    manifest = {
        "title": (
            "AI-assisted decision journaling from public inquiry transcripts: "
            "a UK COVID-19 Inquiry case study"
        ),
        "author": "Akeeb Idowu Lawal",
        "wave": "wave7a_title_page_september",
        "branch": "distinction/final-submission-freeze",
        "parent_commit": "c1feba145bdc46164b5daae79ec1a0d3901e97d0",
        "submission_month": "September 2026",
        "previous_title_month": "May 2026",
        "physical_pages": physical_pages,
        "displayed_pages": disp_pages,
        "body_words": disp_words,
        "markdown_body_words": md_words,
        "word_count_match": disp_words == md_words,
        "docx": docx_out.name,
        "pdf": pdf_path.name,
        "docx_sha256": docx_hash,
        "pdf_sha256": pdf_hash,
        "pdf_mb": round(pdf_path.stat().st_size / (1024 * 1024), 3),
        "supersedes_title_month_of": (
            "outputs/dissertation_integration/run_20260729_153931_wave2_final_integrity_fixes"
        ),
        "created_at_utc": datetime.now(timezone.utc).isoformat(),
    }
    (pkg / "FINAL_SUBMISSION_MANIFEST.json").write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    (pkg / "FINAL_SHA256SUMS.txt").write_text(
        f"{docx_hash}  {docx_out.name}\n{pdf_hash}  {pdf_path.name}\n",
        encoding="utf-8",
    )
    (pkg / "SUBMISSION_PACKAGE_README.md").write_text(
        f"""# Wave 7A title-page month — September 2026

Active formal dissertation candidates after title-month gate.

- Submission month: **September 2026** (was May 2026 on title page only)
- DOCX SHA-256: `{docx_hash}`
- PDF SHA-256: `{pdf_hash}`
- Pages: {physical_pages} physical / {disp_pages} displayed
- Body words: {disp_words:,} (markdown match: {md_words:,})

Wave 2 package `run_20260729_153931_wave2_final_integrity_fixes` remains historical for May-title integrity freeze; superseded for title-page month only.
""",
        encoding="utf-8",
    )

    # pointer for freeze tooling
    pointer = ROOT / "outputs" / "dissertation_integration" / "ACTIVE_FORMAL_SUBMISSION_POINTER.json"
    pointer.write_text(
        json.dumps(
            {
                "active_package": pkg.relative_to(ROOT).as_posix(),
                "docx": (pkg / docx_out.name).relative_to(ROOT).as_posix(),
                "pdf": (pkg / pdf_path.name).relative_to(ROOT).as_posix(),
                "docx_sha256": docx_hash,
                "pdf_sha256": pdf_hash,
                "submission_month": "September 2026",
                "updated_at_utc": datetime.now(timezone.utc).isoformat(),
            },
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )

    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()
