#!/usr/bin/env python3
import zipfile
from pathlib import Path
from docx import Document

path = Path(r"c:\SURREY\MODULES\SEMESTER 2\MSC PROJECT\code\dissertation\Lawal_MSc_Dissertation_submission.docx")
doc = Document(path)
print(f"Size: {path.stat().st_size // 1024} KB, paragraphs: {len(doc.paragraphs)}")

LEAKS = [
    "Joint External Evaluation", "Jesutomiwa", "Kanojia",
    "<Technical CHAPTER>", "<Section title>", "Example figure",
    "Test</w:t></w:r></w:p></w:tc>",  # sample table cell
]
with zipfile.ZipFile(path) as z:
    media = [n for n in z.namelist() if n.startswith("word/media/")]
    print(f"Media: {len(media)}")
    blob = "\n".join(z.read(n).decode("utf-8", errors="ignore") for n in z.namelist() if n.endswith(".xml"))
    for term in LEAKS:
        print(f"  leak {term[:40]!r}: {term in blob}")

visible_opening = any(p.text.strip() == "Opening" for p in doc.paragraphs)
print(f"  visible Opening heading: {visible_opening}")

toc_snap = sum(1 for p in doc.paragraphs if p.style.name.startswith("toc "))
print(f"  toc snapshot paras: {toc_snap}")

for i, p in enumerate(doc.paragraphs):
    if "Chapter 1" in p.text:
        print(f"Chapter 1 at {i}: {p.text[:60]}")
        break

for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Abstract":
        print("Abstract block:")
        for j in range(i, min(i + 6, len(doc.paragraphs))):
            print(f"  {j} [{doc.paragraphs[j].style.name}]: {doc.paragraphs[j].text[:70]}")
        break

for i, p in enumerate(doc.paragraphs):
    if "Table of Contents" in p.text and p.style.name == "Unnumbered 1":
        print("TOC block:")
        for j in range(i, min(i + 4, len(doc.paragraphs))):
            print(f"  {j} [{doc.paragraphs[j].style.name}]: {doc.paragraphs[j].text[:50] or '(field)'}")
        break

figure_caps = [p.text[:60] for p in doc.paragraphs if p.text.startswith("Figure 4.") or p.text.startswith("Figure 3.")]
print(f"Figure captions: {len(figure_caps)}")
for c in figure_caps:
    print(f"  {c}")
