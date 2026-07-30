"""Leak-term scanner: legitimate JEE phrasing passes; genuine leaks still fail."""
from __future__ import annotations

import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DISS = ROOT / "dissertation"
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))
from active_formal_submission import active_docx  # noqa: E402
from build_submission_docx import LEAK_TERMS, verify_output  # noqa: E402

ACTIVE_DOCX = active_docx()


def _docx_with_visible_text(tmp_path: Path, text: str) -> Path:
    """Write a valid DOCX whose visible paragraph contains ``text`` (XML-escaped by python-docx)."""
    path = tmp_path / "probe.docx"
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def _docx_with_raw_xml_substring(tmp_path: Path, raw_substring: str) -> Path:
    """Inject ``raw_substring`` into a sidecar XML part (CDATA) so angle brackets stay intact.

    ``verify_output`` scans all ``*.xml`` zip members as text. Genuine angle-bracket
    template leftovers appear unescaped in OOXML; CDATA reproduces that for probes.
    """
    path = tmp_path / "probe_raw.docx"
    Document().save(path)
    with zipfile.ZipFile(path, "a") as z:
        z.writestr(
            "word/leak_probe.xml",
            f"<probe><![CDATA[{raw_substring}]]></probe>",
        )
    return path


def test_joint_external_evaluation_not_in_leak_terms() -> None:
    assert "Joint External Evaluation" not in LEAK_TERMS
    assert "JEE" not in LEAK_TERMS


def test_genuine_leak_terms_still_configured() -> None:
    for term in (
        "Jesutomiwa",
        "Kanojia",
        "<Technical CHAPTER>",
        "<Section title>",
        "Example figure and caption",
    ):
        assert term in LEAK_TERMS


def test_legitimate_jee_phrase_does_not_trigger_leak_scan(tmp_path: Path) -> None:
    path = _docx_with_visible_text(
        tmp_path,
        "The preparedness mapping used the Joint External Evaluation (JEE) tool.",
    )
    issues = verify_output(path)
    assert "Joint External Evaluation" not in issues
    assert "JEE" not in issues


@pytest.mark.parametrize(
    "leak",
    [
        "Jesutomiwa",
        "Kanojia",
        "Example figure and caption",
    ],
)
def test_genuine_visible_leak_terms_still_detected(tmp_path: Path, leak: str) -> None:
    path = _docx_with_visible_text(tmp_path, f"Draft residue: {leak}")
    issues = verify_output(path)
    assert leak in issues


@pytest.mark.parametrize(
    "leak",
    [
        "<Technical CHAPTER>",
        "<Section title>",
        "<w:t>Test</w:t>",
        "<w:t>Testing</w:t>",
    ],
)
def test_genuine_xml_template_leak_terms_still_detected(tmp_path: Path, leak: str) -> None:
    path = _docx_with_raw_xml_substring(tmp_path, f"Draft residue: {leak}")
    issues = verify_output(path)
    assert leak in issues


def test_active_candidate_docx_jee_not_flagged_when_present() -> None:
    if not ACTIVE_DOCX.is_file():
        pytest.skip("Active formal DOCX missing")
    issues = verify_output(ACTIVE_DOCX)
    assert "Joint External Evaluation" not in issues
    assert "Jesutomiwa" not in issues
    assert "Kanojia" not in issues


def test_dissertation_jee_expansion_unchanged() -> None:
    """Lock intentional JEE / Joint External Evaluation usage; do not rewrite prose."""
    ch2 = (DISS / "CHAPTER_2_LITERATURE.md").read_text(encoding="utf-8")
    ch3 = (DISS / "CHAPTER_3_METHODS.md").read_text(encoding="utf-8")
    abstract = (DISS / "ABSTRACT.md").read_text(encoding="utf-8")
    ch4 = (DISS / "CHAPTER_4_RESULTS.md").read_text(encoding="utf-8")
    ch5 = (DISS / "CHAPTER_5_DISCUSSION.md").read_text(encoding="utf-8")

    assert "Joint External Evaluation" in abstract
    assert "Joint External Evaluation" in ch2
    assert "Joint External Evaluation (JEE)" in ch3
    assert "Joint External Evaluation (JEE)" in ch4
    assert "JEE" in ch3
    assert "JEE" in ch5
